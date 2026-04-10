import io

from docx import Document


def extract_docx_text(docx_bytes: bytes) -> str:
    """Returns plain text including headers, body, and tables."""
    doc = Document(io.BytesIO(docx_bytes))
    parts = []

    # Headers (per-section — contact info often lives here)
    for section in doc.sections:
        for para in section.header.paragraphs:
            if para.text.strip():
                parts.append(para.text)

    # Body paragraphs
    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text)

    # Tables (many resumes use table layout)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    parts.append(cell.text)

    return "\n".join(parts)
