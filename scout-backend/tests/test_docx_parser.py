import io
import pytest
from docx import Document
from app.services.docx_parser import extract_docx_text


def make_simple_docx(body_text: str = "Software Engineer at Acme Corp") -> bytes:
    doc = Document()
    doc.add_paragraph(body_text)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def make_docx_with_table(cell_text: str = "Python, FastAPI") -> bytes:
    doc = Document()
    table = doc.add_table(rows=1, cols=2)
    table.cell(0, 0).text = cell_text
    table.cell(0, 1).text = "Skills"
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


class TestExtractDocxText:
    def test_extracts_body_paragraphs(self):
        docx_bytes = make_simple_docx("Software Engineer at Acme Corp")
        result = extract_docx_text(docx_bytes)
        assert "Software Engineer" in result

    def test_extracts_table_cells(self):
        docx_bytes = make_docx_with_table("Python, FastAPI")
        result = extract_docx_text(docx_bytes)
        assert "Python" in result
